"""Apply Word track-changes (w:ins / w:del) to DOCX contracts."""

from __future__ import annotations

import zipfile
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable
from uuid import uuid4

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from ..config import TRACK_CHANGES_AUTHOR
from ..models import ChangeType, ContractChange

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _revision_id() -> str:
    return str(uuid4().int % 10_000_000)


def _now_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _make_ins(text: str, author: str) -> OxmlElement:
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), _revision_id())
    ins.set(qn("w:author"), author)
    ins.set(qn("w:date"), _now_iso())
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    run.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    ins.append(run)
    return ins


def _make_del(text: str, author: str) -> OxmlElement:
    dele = OxmlElement("w:del")
    dele.set(qn("w:id"), _revision_id())
    dele.set(qn("w:author"), author)
    dele.set(qn("w:date"), _now_iso())
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    run.append(rpr)
    dt = OxmlElement("w:delText")
    dt.set(qn("xml:space"), "preserve")
    dt.text = text
    run.append(dt)
    dele.append(run)
    return dele


def enable_track_revisions(docx_path: Path) -> None:
    """Ensure word/settings.xml contains <w:trackRevisions/>."""
    tmp = docx_path.with_suffix(".tracktmp.docx")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/settings.xml":
                root = etree.fromstring(data)
                ns = {"w": W_NS}
                if root.find("w:trackRevisions", namespaces=ns) is None:
                    node = etree.Element(f"{{{W_NS}}}trackRevisions")
                    root.insert(0, node)
                    data = etree.tostring(
                        root,
                        xml_declaration=True,
                        encoding="UTF-8",
                        standalone=True,
                    )
            zout.writestr(item, data)
    tmp.replace(docx_path)


def _clear_paragraph_runs(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag in {qn("w:r"), qn("w:ins"), qn("w:del"), qn("w:hyperlink")}:
            p.remove(child)


def _paragraph_matches(paragraph, needle: str) -> bool:
    if not needle:
        return False
    text = paragraph.text or ""
    return needle.strip().lower() in text.lower()


def apply_tracked_changes(
    source_docx: str | Path,
    output_docx: str | Path,
    changes: Iterable[ContractChange | dict],
    *,
    author: str = TRACK_CHANGES_AUTHOR,
) -> list[ContractChange]:
    """Apply additions/removals/modifications as Word revision marks.

    Returns the normalized list of changes that were applied (best-effort match
    against paragraph text). Unmatched changes are still recorded with
    metadata.matched=false and appended as a tracked note paragraph.
    """
    source = Path(source_docx)
    output = Path(output_docx)
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(source))
    applied: list[ContractChange] = []

    normalized: list[ContractChange] = []
    for raw in changes:
        if isinstance(raw, ContractChange):
            normalized.append(raw)
        else:
            normalized.append(ContractChange.model_validate(raw))

    for change in normalized:
        matched = False
        for paragraph in doc.paragraphs:
            haystack = change.old_text or change.section
            if not _paragraph_matches(paragraph, haystack):
                continue
            matched = True
            original = paragraph.text
            _clear_paragraph_runs(paragraph)

            if change.change_type == ChangeType.REMOVED:
                paragraph._p.append(_make_del(original or change.old_text, author))
            elif change.change_type == ChangeType.ADDED:
                # Keep original (if any) and insert new text as tracked insertion.
                if original:
                    run = paragraph.add_run(original + "\n")
                    _ = run
                paragraph._p.append(_make_ins(change.new_text or "", author))
            else:  # MODIFIED
                if original or change.old_text:
                    paragraph._p.append(_make_del(original or change.old_text, author))
                if change.new_text:
                    paragraph._p.append(_make_ins(change.new_text, author))
            change.metadata = {**change.metadata, "matched": True, "matched_text": original}
            break

        if not matched:
            # Append a tracked insertion note so legal still sees the proposed edit.
            note = (
                f"[Proposed {change.change_type.value} — {change.section or 'General'}] "
                f"Remove: {change.old_text[:200]} | Add: {change.new_text[:200]} "
                f"| Rationale: {change.rationale[:200]}"
            )
            p = doc.add_paragraph()
            p._p.append(_make_ins(note, author))
            change.metadata = {**change.metadata, "matched": False, "appended_note": True}

        applied.append(change)

    doc.save(str(output))
    enable_track_revisions(output)
    return applied


def ensure_docx_from_text(text: str, output_docx: str | Path) -> Path:
    """Create a simple DOCX from plain text (for fixtures / text uploads)."""
    output = Path(output_docx)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for block in text.split("\n"):
        doc.add_paragraph(block)
    doc.save(str(output))
    enable_track_revisions(output)
    return output
